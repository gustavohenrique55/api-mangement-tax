from pathlib import Path
from datetime import date

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(r"C:\Users\gusta\Documents\ChatGPT\API - MANAGEMNET tAX")
OUT = ROOT / "deliverables" / "Relatorio_Validacao_Juridica_Management_Tax.docx"

GREEN = "147353"
DARK = "14231D"
LIME = "BDDC53"
PAPER = "F3F1E9"
MUTED = "68756F"
RED = "B64335"
WHITE = "FFFFFF"


def shade(cell, color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), color)


def margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cell_text(cell, value, bold=False, color=DARK, size=9):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(value)
    r.bold = bold
    r.font.name = "Calibri"
    r.font.size = Pt(size)
    r.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    margins(cell)


def table(doc, headers, rows, widths=None, font_size=8.5):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    t.style = "Table Grid"
    hdr = t.rows[0]
    set_repeat_table_header(hdr)
    for i, header in enumerate(headers):
        set_cell_text(hdr.cells[i], header, True, WHITE, 8.5)
        shade(hdr.cells[i], GREEN)
        if widths:
            hdr.cells[i].width = Inches(widths[i])
    for r_idx, row in enumerate(rows):
        cells = t.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], str(value), False, DARK, font_size)
            if r_idx % 2 == 1:
                shade(cells[i], PAPER)
            if widths:
                cells[i].width = Inches(widths[i])
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return t


def bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.add_run(text)
    return p


def numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.add_run(text)
    return p


def callout(doc, title, body, color=GREEN):
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    c = t.cell(0, 0)
    shade(c, color)
    margins(c, 140, 180, 140, 180)
    p = c.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = RGBColor(255, 255, 255)
    r.font.size = Pt(11)
    p2 = c.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(body)
    r2.font.color.rgb = RGBColor(255, 255, 255)
    r2.font.size = Pt(10)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def page_break(doc):
    doc.add_page_break()


doc = Document()
sec = doc.sections[0]
sec.page_width = Inches(8.5)
sec.page_height = Inches(11)
sec.top_margin = sec.bottom_margin = Inches(0.78)
sec.left_margin = sec.right_margin = Inches(0.85)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(10.5)
normal.font.color.rgb = RGBColor.from_string(DARK)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.08
for name, size, color in (("Title", 28, DARK), ("Heading 1", 17, GREEN), ("Heading 2", 13, DARK), ("Heading 3", 11, GREEN)):
    st = styles[name]
    st.font.name = "Calibri"
    st.font.size = Pt(size)
    st.font.bold = True
    st.font.color.rgb = RGBColor.from_string(color)
    st.paragraph_format.space_before = Pt(8)
    st.paragraph_format.space_after = Pt(6)

# Cover
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(70)
r = p.add_run("MANAGEMENT TAX\nLATAM & CARIBE")
r.bold = True
r.font.size = Pt(31)
r.font.color.rgb = RGBColor.from_string(DARK)
p2 = doc.add_paragraph()
r = p2.add_run("Pacote de validação jurídica, tributária e de management")
r.bold = True
r.font.size = Pt(18)
r.font.color.rgb = RGBColor.from_string(GREEN)
doc.add_paragraph("Empresa do setor logístico · identificação anonimizada")
doc.add_paragraph("Versão 0.1 · dados e cenários exclusivamente sintéticos")
doc.add_paragraph("Preparado para discussão com advogado tributarista experiente em gestão regional")
doc.add_paragraph().paragraph_format.space_after = Pt(140)
callout(doc, "OBJETIVO DA REUNIÃO", "Validar o desenho de governança, os campos obrigatórios por jurisdição, os limites de decisão do CFO/Board e o protocolo de confiança técnica nos escritórios locais.", DARK)
p = doc.add_paragraph(f"Data de emissão: {date.today().strftime('%d/%m/%Y')}")
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

page_break(doc)
doc.add_heading("1. Resumo executivo", 1)
doc.add_paragraph("A API Management Tax transforma o plano de ação em uma camada de governança regional. Ela organiza dados, alçadas, evidências, desempenho, riscos e oportunidades; não substitui a análise jurídica do profissional habilitado em cada país ou território.")
table(doc, ["Pergunta executiva", "Resposta do sistema"], [
    ("Onde há geração de valor?", "ETR, incentivos, créditos, economia e oportunidades de recuperação."),
    ("Onde há exposição?", "Demandas, contingências, aging, materialidade e posições fora do apetite."),
    ("Quem responde tecnicamente?", "Escritório local habilitado, com escopo, SLA, procuração e reliance documentados."),
    ("Quem decide?", "Alçadas do escritório, Management Tax, CFO ou Board conforme materialidade e risco."),
    ("O que está atualizado?", "Radar legal com fonte, vigência, data de corte, impacto, responsável e validação local."),
], [2.1, 4.7], 9)
callout(doc, "LIMITAÇÃO ESSENCIAL", "Nenhum dado do protótipo constitui parecer, cálculo fiscal definitivo ou confirmação da operação real. A legislação, jurisprudência, súmulas e atos administrativos devem ser obtidos de fontes oficiais/licenciadas e validados pelo assessor local.", RED)
doc.add_heading("Condições anteriores ao piloto", 2)
for item in [
    "Calibrar pesos, metas e limiares com dados históricos; até lá, o scorecard é baseline e não autoriza panel review.",
    "Validar tecnicamente os campos, fontes e obrigações de cada jurisdição com o escritório local.",
    "Medir a carga real de trabalho e confirmar se o modelo cabe na capacidade de gestão disponível.",
    "Obter decisões formais do CFO/Board sobre apetite de risco, materialidade, prioridades e estrutura de custos.",
]: bullet(doc, item)

page_break(doc)
doc.add_heading("2. O que foi incorporado à API", 1)
modules = [
    ("1", "Presença fiscal e operação logística", "Entidades, estabelecimentos, rotas, armazéns, agentes, EP e perfil operacional."),
    ("2", "Compliance e calendário", "Demandas, obrigações, prazos, responsáveis, SLA, documentos e evidências."),
    ("3", "Aduaneiro e comércio exterior", "Regimes especiais, drawback/RECOF/IMMEX, zonas francas, importação e exportação."),
    ("4", "Tributos indiretos", "IVA/VAT/ICMS e equivalentes, créditos, recuperação e impacto em frete/armazenagem."),
    ("5", "Tributos diretos e ETR", "Medições de ETR, reconciliação, incentivos, perdas e planejamento."),
    ("6", "Preços de transferência e intercompany", "Rotas, serviços, royalties, financiamento, beneficiário efetivo e substância."),
    ("7", "Contencioso e contingências", "Probabilidade, exposição, aging, provisão, garantias e estratégia."),
    ("8", "Recuperação de tributos", "Oportunidade, tese, período, valor estimado/validado/realizado e gate jurídico."),
    ("9", "Rede de escritórios", "Cadastro, escopo, honorários, SLA, procuração, privilege, cadence e scorecard."),
    ("10", "Governança executiva", "Configuração do CFO, materialidade, apetite, pesos por país, alçadas, segunda opinião e reliance."),
    ("11", "Radar legal e conhecimento", "Regras, fontes, vigência, alertas, impacto e validação; lições aprendidas."),
    ("12", "Auditoria e segurança", "Tenant isolation, RLS, correlação, RBAC sintético e trilha de auditoria."),
]
table(doc, ["#", "Módulo", "Aplicação"], modules, [0.35, 2.15, 4.3], 8.3)

page_break(doc)
doc.add_heading("3. Cobertura por blocos e natureza jurídica", 1)
doc.add_paragraph("O catálogo atual contém 40 jurisdições: 33 Estados soberanos e 7 territórios/jurisdições não soberanas. Presença operacional, tier definitivo e contratação de assessoria continuam pendentes de confirmação.")
table(doc, ["Bloco de gestão", "Estados", "Territórios", "Total"], [
    ("América Central*", "8", "0", "8"),
    ("Ilhas do Caribe", "13", "6", "19"),
    ("América do Sul", "12", "1", "13"),
    ("Total", "33", "7", "40"),
], [3.5, 1.0, 1.1, 1.0], 9)
doc.add_paragraph("* O México integra o bloco apenas para gestão do portfólio; isso não altera sua classificação geográfica.")
doc.add_heading("Territórios modelados separadamente", 2)
table(doc, ["Jurisdição", "Vínculo soberano", "Modelo jurídico a validar"], [
    ("Porto Rico", "Estados Unidos", "Sistema local próprio e coordenação com regras federais dos EUA."),
    ("Aruba", "Reino dos Países Baixos", "Legislação local e matérias do Reino."),
    ("Curaçao", "Reino dos Países Baixos", "Legislação local e matérias do Reino."),
    ("Ilhas Cayman", "Reino Unido", "Autoridade/lei local no vínculo constitucional britânico."),
    ("Guadalupe", "França", "Direito francês, União Europeia e regras locais aplicáveis."),
    ("Martinica", "França", "Direito francês, União Europeia e regras locais aplicáveis."),
    ("Guiana Francesa", "França", "Direito francês, União Europeia e regras locais aplicáveis."),
], [1.45, 1.75, 3.8], 8.4)
callout(doc, "REGRA DE MODELAGEM", "Território não herda automaticamente toda a legislação do soberano. A API registra o vínculo, mas exige validação simultânea do regime local, do marco soberano e, quando aplicável, do direito da União Europeia.", DARK)

page_break(doc)
doc.add_heading("4. Modelo de supervisão por tiers", 1)
table(doc, ["Tier", "Referência inicial", "Intensidade de gestão", "Pontos especiais"], [
    ("1", "Brasil, México, Colômbia", "Próxima; scorecard trimestral e gates", "Alta materialidade/complexidade."),
    ("2", "Argentina, Chile, Peru, Equador, Paraguai, Uruguai, Bolívia", "Periódica", "Argentina recebe tratamento cambial equivalente a Tier 1."),
    ("3", "América Central", "Hub regional", "CAUCA/RECAUCA onde aplicável; Panamá separado."),
    ("4", "Caribe e Guianas", "Por tradição jurídica, substância e exceção", "Territórios, common law/civil law e baixa tributação."),
    ("5", "Venezuela, Cuba, Haiti", "Reforçada", "Sanções, câmbio e continuidade operacional."),
], [0.45, 2.2, 1.8, 2.55], 8)
doc.add_heading("Quatro controles compensatórios", 2)
for title, body in [
    ("Alçadas formais", "Definem quem pode decidir por faixa de materialidade."),
    ("Gate ex ante", "Exige aprovação antes de posição, disputa ou renúncia relevante."),
    ("Segunda opinião", "Reduz concentração de risco em parecer único."),
    ("Reliance documentado", "Registra questão, parecer, profissional, data, premissas e decisão."),
]:
    p = doc.add_paragraph()
    r = p.add_run(title + " — ")
    r.bold = True
    p.add_run(body)
doc.add_heading("Matriz de alçadas a preencher", 2)
table(doc, ["Evento", "Escritório", "Management Tax", "CFO", "Board"], [
    ("Rotina abaixo do limiar", "Executa", "Monitora", "—", "—"),
    ("Posição material", "Recomenda", "Gate", "Aprova*", "—"),
    ("Fora do apetite", "Parecer", "Escala", "Recomenda", "Decide"),
    ("Risco reputacional/sanções", "Parecer", "Escala imediata", "Decide/escala", "Conforme política"),
], [1.65, 1.25, 1.45, 1.15, 1.25], 8.2)
doc.add_paragraph("* Faixas e valores dependem de decisão formal. A tabela descreve o fluxo, não fixa alçadas financeiras.")

page_break(doc)
doc.add_heading("5. Indicadores e regras de uso", 1)
doc.add_paragraph("O IDE transforma processo e resultado em evidência comparável. Os pesos abaixo são baseline do plano e precisam de calibração antes de produzirem consequência contratual ou panel review.")
table(doc, ["Dimensão", "Indicador", "Peso baseline"], [
    ("Qualidade", "Êxito em disputas", "15%"),
    ("Qualidade", "Retrabalho de pareceres", "10%"),
    ("Prazo", "Tempo de resposta", "10%"),
    ("Prazo", "Cumprimento de prazos legais", "10%"),
    ("Resultado", "Créditos identificados/realizados", "15%"),
    ("Resultado", "Gestão de contingências", "10%"),
    ("Resultado", "Economia tributária", "5%"),
    ("Governança", "Custo-benefício", "10%"),
    ("Governança", "Falhas de processo/compliance", "5%"),
    ("Governança", "Aderência ao apetite", "5%"),
    ("Governança", "Alertas e radar legal", "5%"),
], [1.25, 4.15, 1.25], 8.5)
doc.add_paragraph("Fórmula: IDE = soma de cada indicador normalizado multiplicado por seu peso. A API mantém também ETR, créditos, exposição, aging, SLA e posições fora do apetite no dashboard executivo.")
callout(doc, "PERGUNTA AO ESPECIALISTA", "O scorecard mede desempenho operacional sem sugerir que o Management Tax revisa ou certifica o mérito jurídico do parecer local? Quais indicadores devem variar entre civil law, common law e territórios?", GREEN)

page_break(doc)
doc.add_heading("6. Recuperação de tributos e geração de valor", 1)
doc.add_paragraph("A API possui item próprio para recuperação tributária. O fluxo separa hipótese, validação jurídica, valor economicamente recuperável e realização financeira, evitando apresentar oportunidade bruta como caixa efetivo.")
table(doc, ["Etapa", "Evidência mínima", "Responsável primário", "Gate"], [
    ("Identificação", "Tributo, período, operação, base e fonte", "Escritório/local tax", "Admissibilidade"),
    ("Qualificação", "Tese, prescrição, documentação e riscos", "Escritório local", "Parecer"),
    ("Valoração", "Estimado, validado, custo e prazo", "Tax + Finance", "Materialidade"),
    ("Aprovação", "Apetite, estratégia e alçada", "Management Tax/CFO", "Ex ante"),
    ("Execução", "Pedido, compensação ou restituição", "Escritório local", "Protocolo"),
    ("Realização", "Crédito homologado e efeito caixa", "Finance/Tax", "Reconciliação"),
], [1.1, 2.55, 1.65, 1.3], 8.3)
doc.add_heading("Critérios específicos de logística", 2)
for item in [
    "Frete, armazenagem, desembaraço, handling, combustível e ativos operacionais.",
    "Importação/exportação, regimes aduaneiros especiais e zonas francas.",
    "Créditos de tributos indiretos e restrições por natureza da despesa.",
    "Rotas intercompany, serviços compartilhados, royalties e financiamento de CAPEX.",
    "Nexo, estabelecimento permanente, agentes dependentes e presença continuada.",
]: bullet(doc, item)
callout(doc, "NÃO CONFUNDIR", "Valor identificado ≠ valor validado ≠ valor protocolado ≠ valor homologado ≠ caixa realizado. O dashboard deve mostrar as etapas separadamente.", DARK)

page_break(doc)
doc.add_heading("7. Radar legal, jurisprudência e atualização", 1)
doc.add_paragraph("A API está preparada para registrar e governar atualizações, mas não possui acesso automático universal a todas as leis, súmulas e acórdãos. O acesso depende de conectores com fontes oficiais, bases licenciadas e entregas dos escritórios locais.")
table(doc, ["Camada", "Origem recomendada", "Campos de controle"], [
    ("Legislação", "Diários oficiais, autoridades fiscais, parlamentos e gazetas", "Fonte, ato, publicação, vigência, revogação, jurisdição."),
    ("Administração", "Autoridade fiscal/aduaneira", "Súmula, ruling, instrução, solução, efeito vinculante."),
    ("Jurisprudência", "Tribunais e repositórios oficiais/licenciados", "Órgão, processo, data, tese, status, precedência."),
    ("Doutrina/alerta", "Escritório local e base especializada", "Autor, data, premissas, abrangência, revisão local."),
    ("Impacto interno", "Management Tax + operação", "ETR, exposição, processos, sistemas, prazo e responsável."),
], [1.15, 2.35, 3.15], 8.2)
doc.add_heading("Workflow mínimo de atualização", 2)
for item in [
    "Capturar a fonte e preservar versão/hash ou link verificável.",
    "Classificar jurisdição, tributo, tipo de ato, vigência e revogações.",
    "Submeter relevância e interpretação ao assessor local habilitado.",
    "Avaliar impacto em ETR, risco, contratos, rotas, preços e compliance.",
    "Aprovar ação, registrar responsável/prazo e manter trilha de decisão.",
    "Revalidar periodicamente; alertas vencidos reduzem o IDE do escritório.",
]: numbered(doc, item)
callout(doc, "PRINCÍPIO", "A automação pode encontrar, comparar e alertar. A conclusão jurídica permanece atribuída ao profissional habilitado, com data de corte e escopo claramente registrados.", GREEN)

page_break(doc)
doc.add_heading("8. Matriz de validação por jurisdição", 1)
doc.add_paragraph("Usar uma linha por combinação jurisdição × tema relevante. A validação deve alcançar a legislação local e, nos territórios, o marco do soberano e/ou da União Europeia quando aplicável.")
table(doc, ["Campo", "Resposta a obter", "Evidência"], [
    ("Presença", "Entidades, estabelecimentos, rotas, agentes e ativos", "Organograma + registro local"),
    ("Tributos", "Diretos, indiretos, folha, aduaneiro e retenções", "Mapa tributário vigente"),
    ("Compliance", "Obrigações, calendário, sistemas e penalidades", "Calendário validado"),
    ("Logística", "Frete, armazenagem, porto, importação/exportação", "Fluxos e contratos"),
    ("Incentivos", "Regimes especiais, requisitos e sunset clauses", "Ato concessório/parecer"),
    ("Recuperação", "Créditos, período, prazo prescricional e procedimento", "Memorando de tese"),
    ("Intercompany", "TP, WHT, tratado, beneficial owner e substância", "Estudo/contratos"),
    ("Contencioso", "Processos, exposição, probabilidade, garantias e aging", "Relatório processual"),
    ("Atualização", "Fontes, responsável e cadência", "Radar legal"),
    ("Privilege", "Confidencialidade e circulação do parecer", "Protocolo local"),
    ("Alçada", "Quem recomenda, aprova e executa", "Matriz aprovada"),
    ("Validador", "Nome, registro profissional, escritório, data e escopo", "Assinatura/reliance"),
], [1.15, 3.65, 1.85], 8.1)
doc.add_heading("Status padronizados", 2)
doc.add_paragraph("NÃO INICIADO · EM COLETA · EM VALIDAÇÃO LOCAL · VALIDADO · VALIDADO COM RESSALVA · EXPIRADO · SUBSTITUÍDO")
doc.add_paragraph("Todo item validado deve possuir data de corte, fonte primária, responsável e próxima revisão. ‘Validado’ sem esses quatro elementos deve ser rejeitado pela API.")

page_break(doc)
doc.add_heading("9. Questionário para o advogado tributarista", 1)
questions = [
    "A arquitetura separa corretamente governança regional de execução técnica local?",
    "Quais decisões não podem ser delegadas ao Management Tax ou ao escritório?",
    "Como formalizar privilege/confidencialidade em cada tradição jurídica relevante?",
    "Quais posições exigem obrigatoriamente segunda opinião?",
    "Quais limiares de materialidade devem combinar valor absoluto, EBITDA e reputação?",
    "Quais dados e documentos são indispensáveis para sustentar reliance defensável?",
    "Como tratar territórios com regime próprio e vínculo soberano sem presumir aplicação integral da lei do soberano?",
    "Quais países exigem representante legal, procuração ou registro local específico?",
    "Quais fontes oficiais/licenciadas são adequadas para legislação e jurisprudência em cada país?",
    "Qual periodicidade mínima do radar legal por tier e tema?",
    "Quais métricas do IDE podem gerar incentivo perverso ou invadir o mérito jurídico?",
    "Como separar contingências herdadas das geradas após a data de corte?",
    "Quais regras de prescrição/decadência devem bloquear oportunidades de recuperação?",
    "Como documentar DAC6, Pilar 2/GloBE, CFC, substância e sanções nos gates ex ante?",
    "O fluxo de aprovação atende auditoria, Board e seguradoras sem expor pareceres além do necessário?",
]
for index, q in enumerate(questions, start=1):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.22)
    p.paragraph_format.first_line_indent = Inches(-0.22)
    p.add_run(f"{index}.  {q}")
doc.add_heading("Entregável esperado da validação", 2)
doc.add_paragraph("Ata de validação com: parâmetros aprovados; lacunas por jurisdição; campos obrigatórios; fontes autorizadas; ressalvas; responsáveis; datas-alvo; e autorização ou não para o piloto.")

page_break(doc)
doc.add_heading("10. Roteiro de apresentação e decisões", 1)
table(doc, ["Tempo", "Tema", "Decisão buscada"], [
    ("0–10 min", "Mandato, limites e arquitetura", "Concordância com o papel regional."),
    ("10–25 min", "Tiers, territórios e rede de escritórios", "Ajustes de supervisão e reliance."),
    ("25–40 min", "Indicadores, recuperação e risco", "Validade das métricas e gates."),
    ("40–55 min", "Radar legal e matriz país", "Fontes, frequência e campos obrigatórios."),
    ("55–60 min", "Próximos passos", "Responsáveis, piloto e cronograma."),
], [0.9, 2.6, 3.15], 8.7)
doc.add_heading("Decisões reservadas", 2)
table(doc, ["Decisão", "Proponente", "Aprovador"], [
    ("Apetite de risco e exceções", "Management Tax", "CFO/Board"),
    ("Materialidade e alçadas", "Management Tax + Finance", "CFO/Board"),
    ("Prioridade/weights por país", "Management Tax", "CFO"),
    ("Estrutura de custos e panel", "Management Tax/Procurement", "CFO"),
    ("Mérito jurídico local", "Escritório habilitado", "Conforme lei/alçada"),
    ("Piloto e uso de dados reais", "Management Tax + Jurídico + TI", "Sponsor competente"),
], [3.0, 1.8, 1.85], 8.5)
callout(doc, "RESULTADO DA REUNIÃO", "Parâmetros aprovados, lacunas atribuídas, fontes definidas, protocolo de reliance aceito e piloto autorizado — ou pendências claramente documentadas.", DARK)

page_break(doc)
doc.add_heading("11. Sequência de implantação", 1)
table(doc, ["Fase", "Prazo", "Entregas na API"], [
    ("0", "0–30 dias", "Inventário real, presença, baseline, escritórios, procurações e cobertura."),
    ("1", "30–90 dias", "CFO config, tiers, SLA, IDE baseline, comitê e alçadas."),
    ("2", "90–180 dias", "ETR, incentivos, recuperação, intercompany, DAC6/Pilar 2 e caixa."),
    ("3", "90–180+", "Apetite, matriz de riscos, radar legal, TCF, crises e KPIs."),
    ("4", "180–365 dias", "Reporte consolidado, lições aprendidas e cultura regional."),
    ("Paralela", "Desde o dia 1", "Fila de demandas correntes e cumprimento de prazos legais."),
], [0.85, 1.1, 4.8], 8.5)
doc.add_heading("Próximos passos técnicos", 2)
for item in [
    "Homologar PostgreSQL/RLS no ambiente corporativo e integrar identidade real.",
    "Conectar fontes oficiais/licenciadas e definir política de retenção de evidências.",
    "Carregar dados reais somente após aprovação, classificação e controles de acesso.",
    "Executar piloto em Brasil, México e Colômbia; manter Argentina no monitoramento cambial reforçado.",
    "Calibrar o IDE ao fim do primeiro trimestre e documentar qualquer alteração de peso.",
    "Expandir para Tiers 2–5 por gates de prontidão, não por cobertura nominal.",
]: bullet(doc, item)

page_break(doc)
doc.add_heading("12. Termo de validação", 1)
doc.add_paragraph("O presente documento descreve um protótipo de governança e seus requisitos de validação. Não contém aconselhamento jurídico ou tributário e não deve ser usado para tomada de posição perante autoridades sem parecer local competente.")
table(doc, ["Item", "Registro"], [
    ("Escopo revisado", "____________________________________________"),
    ("Ressalvas", "____________________________________________\n____________________________________________"),
    ("Parâmetros aprovados", "____________________________________________"),
    ("Pendências e responsáveis", "____________________________________________\n____________________________________________"),
    ("Piloto", "(  ) autorizado   (  ) autorizado com condições   (  ) não autorizado"),
    ("Próxima revisão", "____ / ____ / ______"),
], [2.0, 4.75], 9)
doc.add_paragraph("Advogado tributarista / escritório: __________________________________________")
doc.add_paragraph("Registro profissional e jurisdição: _________________________________________")
doc.add_paragraph("Assinatura: ____________________________________  Data: ____ / ____ / ______")
doc.add_paragraph()
callout(doc, "DOCUMENTOS DE APOIO", "Apresentação executiva; especificação consolidada da API; pacote Gate 0; documento de segurança/governança; OpenAPI; dashboard sintético; e trilhas de teste.", GREEN)

# Header/footer
for section in doc.sections:
    hp = section.header.paragraphs[0]
    hp.text = "MANAGEMENT TAX LATAM & CARIBE  |  CONFIDENCIAL"
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in hp.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor.from_string(MUTED)
    fp = section.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run("Empresa anonimizada · Dados sintéticos · Não constitui parecer")
    run.font.name = "Calibri"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(MUTED)

doc.core_properties.title = "Pacote de validação jurídica — Management Tax LATAM & Caribe"
doc.core_properties.subject = "Governança tributária regional para empresa de logística"
doc.core_properties.author = "Management Tax LATAM & Caribe"
OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT)
print(OUT)
